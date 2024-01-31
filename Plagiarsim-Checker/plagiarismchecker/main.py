from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import HTMLResponse
from algorithm import main
from algorithm import fileSimilarity
import PyPDF2
import pytesseract
import io
from PIL import Image
from docx import Document

app = FastAPI()

@app.get("/")
async def home():
    return {"Endpoint is working"}

@app.post("/text_similarity/")
async def test(q: str = Form(...)):
    print("request is welcome test")
    print(q)
    if q:
        percent, link = main.findSimilarity(q)
        percent = round(percent, 2)
    print("Output.....................!!!!!!!!", percent, link)
    return {"link": link, "percent": percent}

@app.post("/doc_similarity/")
async def filetest(docfile: UploadFile = File(...)):
    value = ""
    print(docfile.filename)
    if docfile.filename.endswith(".txt"):
        value = (await docfile.read()).decode()

    elif docfile.filename.endswith(".docx"):
        document = Document(io.BytesIO(await docfile.read()))
        for para in document.paragraphs:
            value += para.text

    elif docfile.filename.endswith(".pdf"):
        pdfReader = PyPDF2.PdfFileReader(io.BytesIO(await docfile.read()))
        pageObj = pdfReader.getPage(0)
        value = pageObj.extractText()

    percent, link = main.findSimilarity(value)
    print("Output...................!!!!!!!!", percent, link)
    return {"link": link, "percent": percent}

@app.post("/image_plagiarism_check/")
async def image_plagiarism_check(image: UploadFile = File(...)):
    if image:
        uploaded_image = Image.open(io.BytesIO(await image.read()))
        extracted_text = pytesseract.image_to_string(uploaded_image)
        percent, link = main.findSimilarity(extracted_text)
        percent = round(percent, 2)
        return {"link": link, "percent": percent}
    return {}



@app.post("/two_text_comparison/")
async def twofiletest1(q1: str = Form(...), q2: str = Form(...)):
    print("Submiited text for 1st and 2nd")
    print(q1)
    print(q2)
    if q1 != "" and q2 != "":
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(q1, q2)
    result = round(result, 2)
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!", result)
    return {"result": result}

@app.post("/two_file_comparison/")
async def twofilecompare1(docfile1: UploadFile = File(...), docfile2: UploadFile = File(...)):
    value1 = ""
    value2 = ""
    if docfile1.filename.endswith(".txt") and docfile2.filename.endswith(".txt"):
        value1 = (await docfile1.read()).decode()
        value2 = (await docfile2.read()).decode()

    elif docfile1.filename.endswith(".docx") and docfile2.filename.endswith(".docx"):
        document1 = Document(io.BytesIO(await docfile1.read()))
        document2 = Document(io.BytesIO(await docfile2.read()))
        for para in document1.paragraphs:
            value1 += para.text
        for para in document2.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1, value2)
    print("Output..................!!!!!!!!", result)
    return {"result": result}
